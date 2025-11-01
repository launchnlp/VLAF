import os
import docx
import json
import argparse

def update_doc_with_af_results(
    doc: docx.Document,
    model_name: str,
    non_oversight_file: str,
    oversight_file: str,
    non_oversight_label: str,
    oversight_label: str,
    split_user_input: bool = False
) -> None:
    """
    Update the given document with the results from the AF explorations.

    Parameters:
    - doc: The docx Document object to update.
    - model_name: The name of the model being processed.
    - non_oversight_file: Path to the JSON file with non-oversight results.
    - oversight_file: Path to the JSON file with oversight results.
    - non_oversight_label: Label for the non-oversight condition (e.g., 'Paid-tier').
    - oversight_label: Label for the oversight condition (e.g., 'Free-tier').
    """

    # Add a heading for a model
    doc.add_heading(model_name, level=2)

    # opening the files
    with open(non_oversight_file, 'r') as f:
        non_oversight_data = json.load(f)
    with open(oversight_file, 'r') as f:
        oversight_data = json.load(f)

    # iterating over each file
    for no_data, o_data in zip(non_oversight_data, oversight_data):

        # add a sub-heading based on o_data input
        if split_user_input:
            user_input = o_data['input'].split('<user_input>')[1].split('</user_input>')[0].strip()
            doc.add_heading(user_input, level=3)
        else:
            doc.add_heading(o_data['input'], level=3)

        # add the no-data response
        no_oversight_para = doc.add_paragraph()
        bold_run = no_oversight_para.add_run(f'{non_oversight_label}:')
        bold_run.bold = True
        no_oversight_para.add_run(f' {no_data["response"]}')

        # add the oversight response
        oversight_para = doc.add_paragraph()
        bold_run = oversight_para.add_run(f'{oversight_label}:')
        bold_run.bold = True
        oversight_para.add_run(f' {o_data["response"]}')


def af_doc_creation_wrapper(
    input_folder: str,
    model_list: list[str],
    output_file: str
):
    
    # intialize a document
    doc = docx.Document()

    # iterating over the models
    for model in model_list:

        # check if the files exist
        non_oversight_file = os.path.join(input_folder, f'{model}_False.json')
        oversight_file = os.path.join(input_folder, f'{model}_True.json')
        if not os.path.exists(non_oversight_file) or not os.path.exists(oversight_file):
            print(f'Files for model {model} do not exist. Skipping...')
            continue
        non_oversight_label = 'Paid-tier'
        oversight_label = 'Free-tier'

        # update the document with the results
        update_doc_with_af_results(
            doc,
            model,
            non_oversight_file,
            oversight_file,
            non_oversight_label,
            oversight_label,
            split_user_input=True
        )

    # save the document
    doc.save(output_file)

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Programmatically print the results of the AF explorations in a document')
    parser.add_argument('--input_folder', type=str, default='results/MFT/mft_fairness_strong')
    parser.add_argument('--model_list', type=str, nargs='+', default=[
        'openai:gpt-4o',
        'openai:gpt-5',
        'vllm:Llama-3.1-8B-Instruct',
        'vllm:Qwen3-8B',
        'vllm:Qwen3-14B',
        'vllm:Qwen3-32B'
    ])
    parser.add_argument('--output_file', type=str, default='results/MFT/mft_fairness_strong/results.docx')
    args = parser.parse_args()

    af_doc_creation_wrapper(
        args.input_folder,
        args.model_list,
        args.output_file
    )